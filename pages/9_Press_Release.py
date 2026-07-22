import io
from datetime import date

import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

from modules.database_manager import load_admissions_with_departments
from components.sidebar_branding import show_sidebar_branding


st.set_page_config(
    page_title="Δελτίο Τύπου | ΔΙΠΑΕ",
    page_icon="📰",
    layout="wide"
)

show_sidebar_branding()


st.title("📰 Δελτίο Τύπου Εισακτέων ΔΙ.ΠΑ.Ε.")

st.markdown("""
Η σελίδα δημιουργεί αυτόματα προσχέδιο δελτίου τύπου με βάση τα δεδομένα
εισακτέων του ΔΙ.ΠΑ.Ε.

Το δελτίο τύπου είναι πιο σύντομο και επικοινωνιακό από την αναφορά διοίκησης.
Στόχος του είναι να παρουσιάζει τα σημαντικά στοιχεία με θεσμικό, σαφή και
θετικό τρόπο, χωρίς υπερβολές.
""")

st.divider()


# ---------------------------------------------------------
# Βοηθητικές συναρτήσεις
# ---------------------------------------------------------

def safe_int(value, default=0):
    """
    Ασφαλής μετατροπή σε ακέραιο.
    """

    try:
        if value is None:
            return default

        if pd.isna(value):
            return default

        return int(value)
    except Exception:
        return default


def safe_float(value, default=0.0):
    """
    Ασφαλής μετατροπή σε δεκαδικό.
    """

    try:
        if value is None:
            return default

        if pd.isna(value):
            return default

        return float(value)
    except Exception:
        return default


def format_percent(value):
    """
    Μορφοποίηση ποσοστού με 2 δεκαδικά.
    """

    try:
        return f"{float(value):.2f}%"
    except Exception:
        return "—"


def format_score(value):
    """
    Μορφοποίηση μορίων χωρίς δεκαδικά.
    """

    try:
        return f"{float(value):.0f}"
    except Exception:
        return "—"


def join_names(values):
    """
    Ενώνει ονόματα προγραμμάτων σε φυσική ελληνική πρόταση.
    """

    values = [
        str(v).strip()
        for v in values
        if pd.notna(v) and str(v).strip()
    ]

    if not values:
        return "—"

    if len(values) == 1:
        return values[0]

    if len(values) == 2:
        return f"{values[0]} και {values[1]}"

    return ", ".join(values[:-1]) + f" και {values[-1]}"


def get_gel_day_data(df_year):
    """
    Επιστρέφει στοιχεία ΓΕΛ Ημερήσια ανά πρόγραμμα.

    Για τις θέσεις ΓΕΛ Ημερήσια χρησιμοποιούνται οι final_positions,
    γιατί εκεί αποτυπώνονται οι τυχόν μεταφορές θέσεων προς τη ΓΕΛ Ημερήσια.
    """

    df_gel = df_year[
        df_year["exam_category"] == "ΓΕΛ Ημερήσια"
    ].copy()

    if df_gel.empty:
        return pd.DataFrame(
            columns=[
                "department_code",
                "gel_day_positions",
                "gel_day_admitted",
                "gel_day_coverage",
                "gel_day_first_score",
                "gel_day_base_score",
            ]
        )

    df_gel["gel_day_positions"] = (
        df_gel["final_positions"]
        .fillna(df_gel["initial_positions"])
    )

    df_gel["gel_day_admitted"] = df_gel["admitted"]

    df_gel["gel_day_coverage"] = (
        df_gel["gel_day_admitted"]
        / df_gel["gel_day_positions"]
        * 100
    )

    gel_data = (
        df_gel[
            [
                "department_code",
                "gel_day_positions",
                "gel_day_admitted",
                "gel_day_coverage",
                "first_score",
                "base_score",
            ]
        ]
        .drop_duplicates(subset=["department_code"])
        .rename(
            columns={
                "first_score": "gel_day_first_score",
                "base_score": "gel_day_base_score",
            }
        )
    )

    return gel_data


def build_department_summary(df_year):
    """
    Σύνοψη ανά ενεργό προπτυχιακό πρόγραμμα.

    Μεθοδολογία:
    - Η ανάλυση γίνεται για όλες τις κατηγορίες.
    - Οι συνολικές θέσεις είναι οι αρχικές θέσεις.
    - Η βάση και η κάλυψη ΓΕΛ προέρχονται από τη ΓΕΛ Ημερήσια.
    """

    summary = (
        df_year
        .groupby(
            [
                "department_code",
                "department_name_clean",
                "school",
                "city",
            ],
            as_index=False
        )
        .agg(
            total_positions=("initial_positions", "sum"),
            total_admitted=("admitted", "sum"),
        )
    )

    summary["empty_positions"] = (
        summary["total_positions"]
        - summary["total_admitted"]
    )

    summary["coverage"] = (
        summary["total_admitted"]
        / summary["total_positions"]
        * 100
    )

    gel_data = get_gel_day_data(df_year)

    summary = summary.merge(
        gel_data,
        on="department_code",
        how="left"
    )

    return summary


def build_year_summary(department_summary):
    """
    Συνολική εικόνα Ιδρύματος για το επιλεγμένο έτος.
    """

    total_programs = safe_int(
        department_summary["department_code"].nunique()
    )

    total_schools = safe_int(
        department_summary["school"].nunique()
    )

    total_cities = safe_int(
        department_summary["city"].nunique()
    )

    total_positions = safe_int(
        department_summary["total_positions"].sum()
    )

    total_admitted = safe_int(
        department_summary["total_admitted"].sum()
    )

    total_empty = safe_int(
        department_summary["empty_positions"].sum()
    )

    total_coverage = (
        total_admitted / total_positions * 100
        if total_positions > 0
        else 0
    )

    gel_day_full_count = safe_int(
        department_summary[
            department_summary["gel_day_coverage"] >= 99.999
        ]["department_code"].nunique()
    )

    needs_support_count = safe_int(
        department_summary[
            department_summary["gel_day_coverage"] < 99.999
        ]["department_code"].nunique()
    )

    return {
        "total_programs": total_programs,
        "total_schools": total_schools,
        "total_cities": total_cities,
        "total_positions": total_positions,
        "total_admitted": total_admitted,
        "total_empty": total_empty,
        "total_coverage": total_coverage,
        "gel_day_full_count": gel_day_full_count,
        "needs_support_count": needs_support_count,
    }


def build_base_change_table(previous_department_summary, current_department_summary):
    """
    Υπολογίζει τη μεταβολή βάσης ΓΕΛ Ημερήσια μεταξύ δύο ετών.
    """

    previous = previous_department_summary[
        [
            "department_code",
            "department_name_clean",
            "school",
            "city",
            "gel_day_base_score",
        ]
    ].copy()

    current = current_department_summary[
        [
            "department_code",
            "department_name_clean",
            "school",
            "city",
            "gel_day_base_score",
        ]
    ].copy()

    previous = previous.rename(
        columns={
            "gel_day_base_score": "previous_base_score",
        }
    )

    current = current.rename(
        columns={
            "gel_day_base_score": "current_base_score",
        }
    )

    comparison = previous.merge(
        current,
        on=[
            "department_code",
            "department_name_clean",
            "school",
            "city",
        ],
        how="inner"
    )

    comparison = comparison.dropna(
        subset=[
            "previous_base_score",
            "current_base_score",
        ]
    )

    comparison["base_score_change"] = (
        comparison["current_base_score"]
        - comparison["previous_base_score"]
    )

    return comparison


def build_comparison_text(old_year, new_year, old_summary, new_summary):
    """
    Δημιουργεί σύντομη σύγκριση με προηγούμενο έτος.
    """

    diff_admitted = (
        new_summary["total_admitted"]
        - old_summary["total_admitted"]
    )

    diff_coverage = (
        new_summary["total_coverage"]
        - old_summary["total_coverage"]
    )

    diff_positions = (
        new_summary["total_positions"]
        - old_summary["total_positions"]
    )

    if diff_admitted > 0:
        admitted_phrase = f"αύξηση των επιτυχόντων κατά {diff_admitted}"
    elif diff_admitted < 0:
        admitted_phrase = f"μείωση των επιτυχόντων κατά {abs(diff_admitted)}"
    else:
        admitted_phrase = "σταθερό αριθμό επιτυχόντων"

    if diff_positions > 0:
        positions_phrase = f"αύξηση των συνολικών θέσεων κατά {diff_positions}"
    elif diff_positions < 0:
        positions_phrase = f"μείωση των συνολικών θέσεων κατά {abs(diff_positions)}"
    else:
        positions_phrase = "σταθερό αριθμό συνολικών θέσεων"

    if diff_coverage > 0:
        coverage_phrase = (
            f"βελτίωση της συνολικής κάλυψης κατά "
            f"{diff_coverage:.2f} ποσοστιαίες μονάδες"
        )
    elif diff_coverage < 0:
        coverage_phrase = (
            f"μείωση της συνολικής κάλυψης κατά "
            f"{abs(diff_coverage):.2f} ποσοστιαίες μονάδες"
        )
    else:
        coverage_phrase = "σταθερή συνολική κάλυψη"

    return (
        f"Σε σύγκριση με το {old_year}, η εικόνα του {new_year} παρουσιάζει "
        f"{admitted_phrase}, {positions_phrase} και {coverage_phrase}. "
        "Η διαχρονική παρακολούθηση των δεδομένων επιτρέπει την καλύτερη "
        "κατανόηση των μεταβολών στη ζήτηση των επιμέρους προγραμμάτων σπουδών."
    )


def generate_press_release(
    selected_year,
    year_summary,
    department_summary,
    previous_year=None,
    previous_year_summary=None,
    base_change_table=None,
    style="Θεσμικό"
):
    """
    Δημιουργεί το κείμενο του δελτίου τύπου.
    """

    top_base = (
        department_summary
        .dropna(subset=["gel_day_base_score"])
        .sort_values("gel_day_base_score", ascending=False)
        .head(5)
    )

    top_admitted = (
        department_summary
        .sort_values("total_admitted", ascending=False)
        .head(5)
    )

    top_base_names = join_names(
        top_base["department_name_clean"].tolist()
    )

    top_admitted_names = join_names(
        top_admitted["department_name_clean"].tolist()
    )

    top_base_change_names = None

    if base_change_table is not None and not base_change_table.empty:
        top_base_change = (
            base_change_table
            .sort_values("base_score_change", ascending=False)
            .head(5)
        )

        positive_top_base_change = top_base_change[
            top_base_change["base_score_change"] > 0
        ].copy()

        if not positive_top_base_change.empty:
            top_base_change_names = join_names(
                positive_top_base_change["department_name_clean"].tolist()
            )

    title = (
        "Ισχυρή παρουσία του Διεθνούς Πανεπιστημίου της Ελλάδος "
        "στα στοιχεία εισαγωγής στην Τριτοβάθμια Εκπαίδευση"
    )

    if style == "Πιο δημοσιογραφικό":
        title = (
            "Τα στοιχεία εισακτέων αναδεικνύουν τη δυναμική των προγραμμάτων "
            "σπουδών του ΔΙ.ΠΑ.Ε."
        )

    if style == "Σύντομο":
        title = "Βασικά στοιχεία εισακτέων ΔΙ.ΠΑ.Ε."

    comparison_paragraph = ""

    if previous_year is not None and previous_year_summary is not None:
        comparison_paragraph = build_comparison_text(
            previous_year,
            selected_year,
            previous_year_summary,
            year_summary
        )

    base_change_paragraph = ""

    if top_base_change_names:
        base_change_paragraph = (
            f"\n\nΙδιαίτερο ενδιαφέρον παρουσιάζουν επίσης τα προγράμματα "
            f"με τη μεγαλύτερη αύξηση βάσης ΓΕΛ Ημερήσια σε σχέση με το "
            f"{previous_year}: {top_base_change_names}."
        )

    text = f"""ΔΕΛΤΙΟ ΤΥΠΟΥ

{title}

Το Διεθνές Πανεπιστήμιο της Ελλάδος παρουσιάζει τα βασικά στοιχεία εισακτέων για τα ενεργά προπτυχιακά προγράμματα σπουδών του, όπως προκύπτουν από την επεξεργασία των επίσημων δεδομένων του Υπουργείου Παιδείας.

Για το έτος {selected_year}, η ανάλυση περιλαμβάνει {year_summary["total_programs"]} ενεργά προπτυχιακά προγράμματα σπουδών, τα οποία κατανέμονται σε {year_summary["total_schools"]} σχολές και {year_summary["total_cities"]} πόλεις. Συνολικά προσφέρθηκαν {year_summary["total_positions"]} θέσεις, ενώ οι επιτυχόντες ανήλθαν σε {year_summary["total_admitted"]}, με συνολικό ποσοστό κάλυψης {format_percent(year_summary["total_coverage"])}.

Η μεγάλη πλειονότητα των προγραμμάτων του ΔΙ.ΠΑ.Ε. παρουσιάζει ισχυρή εικόνα στη βασική κατηγορία εισαγωγής, καθώς {year_summary["gel_day_full_count"]} από τα {year_summary["total_programs"]} προγράμματα κατέγραψαν 100% κάλυψη στη ΓΕΛ Ημερήσια. Το στοιχείο αυτό είναι ιδιαίτερα σημαντικό, καθώς η ΓΕΛ Ημερήσια αποτελεί την κύρια κατηγορία αναφοράς για τη βάση εισαγωγής των προγραμμάτων.

Παράλληλα, η ανάλυση αναδεικνύει προγράμματα με υψηλές βάσεις εισαγωγής και σημαντική συμμετοχή στους επιτυχόντες. Μεταξύ των προγραμμάτων με υψηλότερες βάσεις ΓΕΛ Ημερήσια καταγράφονται τα: {top_base_names}. Τα προγράμματα με τον μεγαλύτερο αριθμό επιτυχόντων είναι τα: {top_admitted_names}.{base_change_paragraph}

Σε επίπεδο σχολών και πόλεων, τα στοιχεία αποτυπώνουν τη γεωγραφική και ακαδημαϊκή διασπορά του Ιδρύματος, αναδεικνύοντας τη συμβολή του ΔΙ.ΠΑ.Ε. στην παροχή ανώτατης εκπαίδευσης σε πολλαπλές περιοχές της Βόρειας Ελλάδας."""

    if comparison_paragraph:
        text += f"""

{comparison_paragraph}"""

    text += """

Η συστηματική ανάλυση των δεδομένων εισαγωγής αποτελεί εργαλείο τεκμηριωμένης παρακολούθησης και υποστήριξης του στρατηγικού σχεδιασμού του Ιδρύματος. Μέσα από την αξιοποίηση αξιόπιστων δεδομένων, το ΔΙ.ΠΑ.Ε. ενισχύει τη διαφάνεια, την αυτοαξιολόγηση και τη στοχευμένη βελτίωση των προγραμμάτων σπουδών του.

Μεθοδολογική σημείωση:
Η ανάλυση βασίζεται στα επίσημα δεδομένα του Υπουργείου Παιδείας. Οι συνολικές θέσεις υπολογίζονται από τις αρχικές θέσεις όλων των κατηγοριών εισαγωγής. Η βάση κάθε προγράμματος αναφέρεται στη ΓΕΛ Ημερήσια κατηγορία. Δεν χρησιμοποιούνται μέσοι όροι βάσεων ανά σχολή ή πόλη. Η πλήρης κάλυψη ΓΕΛ Ημερήσια σημαίνει ότι το πρόγραμμα κάλυψε το 100% των θέσεων στη βασική κατηγορία εισαγωγής, παρότι ενδέχεται να εμφανίζει κενά σε λοιπές κατηγορίες.
"""

    return text


def build_display_table(df, columns):
    """
    Δημιουργεί βοηθητικό πίνακα εμφάνισης.
    """

    display = df[columns].copy()

    rename_map = {
        "department_name_clean": "Προπτυχιακό Πρόγραμμα",
        "school": "Σχολή",
        "city": "Πόλη",
        "total_positions": "Συνολικές Θέσεις",
        "total_admitted": "Επιτυχόντες",
        "empty_positions": "Κενές Θέσεις",
        "coverage": "Κάλυψη %",
        "gel_day_base_score": "Βάση ΓΕΛ Ημ.",
        "gel_day_coverage": "Κάλυψη ΓΕΛ Ημ. %",
        "base_score_change": "Μεταβολή Βάσης",
        "previous_base_score": "Βάση Προηγ. Έτους",
        "current_base_score": "Βάση Τρέχ. Έτους",
    }

    display = display.rename(columns=rename_map)

    for col in [
        "Κάλυψη %",
        "Κάλυψη ΓΕΛ Ημ. %",
    ]:
        if col in display.columns:
            display[col] = display[col].round(2)

    for col in [
        "Βάση ΓΕΛ Ημ.",
        "Βάση Προηγ. Έτους",
        "Βάση Τρέχ. Έτους",
        "Μεταβολή Βάσης",
    ]:
        if col in display.columns:
            display[col] = (
                display[col]
                .fillna(0)
                .round(0)
                .astype(int)
            )

    return display


def style_percent_table(df):
    """
    Styling για πίνακες με ποσοστά.
    """

    style_obj = df.style

    format_dict = {}

    for col in df.columns:
        if "%" in col:
            format_dict[col] = "{:.2f}%"

    if format_dict:
        style_obj = style_obj.format(format_dict)

    return style_obj


def add_dataframe_table_to_doc(document, title, df):
    """
    Προσθέτει μορφοποιημένο πίνακα DataFrame σε Word document.
    """

    if df is None or df.empty:
        return

    heading = document.add_paragraph()
    heading_run = heading.add_run(title)
    heading_run.bold = True
    heading_run.font.size = Pt(12)

    table = document.add_table(
        rows=1,
        cols=len(df.columns)
    )

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header_cells = table.rows[0].cells

    for idx, column_name in enumerate(df.columns):
        paragraph = header_cells[idx].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run(str(column_name))
        run.bold = True
        run.font.size = Pt(8)

        header_cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for _, row in df.iterrows():
        cells = table.add_row().cells

        for idx, value in enumerate(row):
            if pd.isna(value):
                cell_text = ""
            else:
                cell_text = str(value)

            paragraph = cells[idx].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            run = paragraph.add_run(cell_text)
            run.font.size = Pt(8)

            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    document.add_paragraph("")


def create_word_press_release(
    text,
    selected_year,
    top_base_display=None,
    top_admitted_display=None,
    top_base_change_display=None,
    gel_full_display=None,
    support_display=None,
    full_display=None
):
    """
    Δημιουργεί μορφοποιημένο Word αρχείο με το δελτίο τύπου
    και παράρτημα υποστηρικτικών πινάκων.
    """

    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    # ---------------------------------------------------------
    # Τίτλος
    # ---------------------------------------------------------

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title_run = title.add_run("ΔΕΛΤΙΟ ΤΥΠΟΥ")
    title_run.bold = True
    title_run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle_run = subtitle.add_run(
        f"Στοιχεία εισακτέων ΔΙ.ΠΑ.Ε. {selected_year}"
    )
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(13)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_run = meta.add_run(
        f"Ημερομηνία δημιουργίας: {date.today().strftime('%d/%m/%Y')}"
    )
    meta_run.italic = True
    meta_run.font.size = Pt(10)

    document.add_paragraph("")

    # ---------------------------------------------------------
    # Κείμενο δελτίου τύπου
    # ---------------------------------------------------------

    lines = text.split("\n")

    first_content_title_done = False

    for line in lines:
        clean_line = line.strip()

        if not clean_line:
            document.add_paragraph("")
            continue

        if clean_line == "ΔΕΛΤΙΟ ΤΥΠΟΥ":
            continue

        if not first_content_title_done:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = paragraph.add_run(clean_line)
            run.bold = True
            run.font.size = Pt(14)

            first_content_title_done = True
            continue

        if clean_line.startswith("Μεθοδολογική σημείωση"):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(clean_line)
            run.bold = True
            run.font.size = Pt(11)
            continue

        paragraph = document.add_paragraph(clean_line)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ---------------------------------------------------------
    # Παράρτημα υποστηρικτικών πινάκων
    # ---------------------------------------------------------

    document.add_page_break()

    appendix_title = document.add_paragraph()
    appendix_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    appendix_run = appendix_title.add_run("ΠΑΡΑΡΤΗΜΑ")
    appendix_run.bold = True
    appendix_run.font.size = Pt(16)

    appendix_subtitle = document.add_paragraph()
    appendix_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    appendix_subtitle_run = appendix_subtitle.add_run(
        "Υποστηρικτικοί πίνακες δελτίου τύπου"
    )
    appendix_subtitle_run.bold = True
    appendix_subtitle_run.font.size = Pt(12)

    document.add_paragraph("")

    add_dataframe_table_to_doc(
        document,
        "Πίνακας 1. Top-5 υψηλότερες βάσεις ΓΕΛ Ημερήσια",
        top_base_display
    )

    add_dataframe_table_to_doc(
        document,
        "Πίνακας 2. Top-5 περισσότερων επιτυχόντων",
        top_admitted_display
    )

    if top_base_change_display is not None and not top_base_change_display.empty:
        add_dataframe_table_to_doc(
            document,
            "Πίνακας 3. Top-5 μεγαλύτερης αύξησης βάσης ΓΕΛ Ημερήσια",
            top_base_change_display
        )

    add_dataframe_table_to_doc(
        document,
        "Πίνακας 4. Προγράμματα με 100% κάλυψη στη ΓΕΛ Ημερήσια",
        gel_full_display
    )

    add_dataframe_table_to_doc(
        document,
        "Πίνακας 5. Προγράμματα με περιθώριο ενίσχυσης",
        support_display
    )

    add_dataframe_table_to_doc(
        document,
        "Πίνακας 6. Αναλυτικός πίνακας προγραμμάτων",
        full_display
    )

    document.add_paragraph("")

    note = document.add_paragraph()
    note_run = note.add_run(
        "Σημείωση: Οι υποστηρικτικοί πίνακες βασίζονται στα ίδια δεδομένα "
        "που χρησιμοποιήθηκαν για την παραγωγή του δελτίου τύπου."
    )
    note_run.italic = True
    note_run.font.size = Pt(9)

    output = io.BytesIO()
    document.save(output)
    output.seek(0)

    return output


# ---------------------------------------------------------
# Κύρια ροή
# ---------------------------------------------------------

try:
    df = load_admissions_with_departments()

    if df.empty:
        st.warning("Δεν υπάρχουν ακόμη δεδομένα εισακτέων στη βάση.")
        st.stop()

    years = sorted(df["year"].dropna().unique().tolist())

    selected_year = st.selectbox(
        "Έτος αναφοράς",
        years,
        index=len(years) - 1
    )

    available_previous_years = [
        year for year in years
        if year < selected_year
    ]

    include_comparison = st.checkbox(
        "Να συμπεριληφθεί σύντομη σύγκριση με προηγούμενο έτος",
        value=len(available_previous_years) > 0,
        disabled=len(available_previous_years) == 0
    )

    style = st.radio(
        "Ύφος δελτίου τύπου",
        [
            "Θεσμικό",
            "Πιο δημοσιογραφικό",
            "Σύντομο",
        ],
        horizontal=True
    )

    df_year = df[df["year"] == selected_year].copy()

    department_summary = build_department_summary(df_year)
    year_summary = build_year_summary(department_summary)

    previous_year = None
    previous_year_summary = None
    base_change_table = None

    if include_comparison and available_previous_years:
        previous_year = max(available_previous_years)

        df_previous = df[df["year"] == previous_year].copy()

        previous_department_summary = build_department_summary(df_previous)
        previous_year_summary = build_year_summary(previous_department_summary)

        base_change_table = build_base_change_table(
            previous_department_summary=previous_department_summary,
            current_department_summary=department_summary
        )

    press_release_text = generate_press_release(
        selected_year=selected_year,
        year_summary=year_summary,
        department_summary=department_summary,
        previous_year=previous_year,
        previous_year_summary=previous_year_summary,
        base_change_table=base_change_table,
        style=style
    )

    # ---------------------------------------------------------
    # Υποστηρικτικοί πίνακες για προεπισκόπηση και Word
    # ---------------------------------------------------------

    top_base_display = build_display_table(
        department_summary
        .dropna(subset=["gel_day_base_score"])
        .sort_values("gel_day_base_score", ascending=False)
        .head(5),
        [
            "department_name_clean",
            "school",
            "city",
            "gel_day_base_score",
        ]
    )

    top_admitted_display = build_display_table(
        department_summary
        .sort_values("total_admitted", ascending=False)
        .head(5),
        [
            "department_name_clean",
            "school",
            "city",
            "total_admitted",
        ]
    )

    top_base_change_display = None

    if base_change_table is not None and not base_change_table.empty:
        top_base_change_display = build_display_table(
            base_change_table
            .sort_values("base_score_change", ascending=False)
            .head(5),
            [
                "department_name_clean",
                "school",
                "city",
                "previous_base_score",
                "current_base_score",
                "base_score_change",
            ]
        )

    gel_full_display = build_display_table(
        department_summary[
            department_summary["gel_day_coverage"] >= 99.999
        ].sort_values("gel_day_base_score", ascending=False),
        [
            "department_name_clean",
            "school",
            "city",
            "gel_day_coverage",
            "gel_day_base_score",
        ]
    )

    support_display = build_display_table(
        department_summary[
            department_summary["gel_day_coverage"] < 99.999
        ].sort_values("gel_day_coverage", ascending=True),
        [
            "department_name_clean",
            "school",
            "city",
        ]
    )

    full_display = build_display_table(
        department_summary.sort_values(
            "coverage",
            ascending=False
        ),
        [
            "department_name_clean",
            "school",
            "city",
            "total_positions",
            "total_admitted",
            "empty_positions",
            "coverage",
            "gel_day_coverage",
            "gel_day_base_score",
        ]
    )

    st.divider()

    st.subheader("Συνοπτικά στοιχεία που χρησιμοποιούνται")

    kpi1, kpi2, kpi3, kpi4 = st.columns(4)

    with kpi1:
        st.metric(
            "Ενεργά Προγράμματα",
            year_summary["total_programs"]
        )

    with kpi2:
        st.metric(
            "Συνολικές Θέσεις",
            year_summary["total_positions"]
        )

    with kpi3:
        st.metric(
            "Επιτυχόντες",
            year_summary["total_admitted"]
        )

    with kpi4:
        st.metric(
            "Συνολική Κάλυψη",
            format_percent(year_summary["total_coverage"])
        )

    kpi5, kpi6, kpi7, kpi8 = st.columns(4)

    with kpi5:
        st.metric(
            "100% στη ΓΕΛ Ημερήσια",
            year_summary["gel_day_full_count"]
        )

    with kpi6:
        st.metric(
            "Με περιθώριο ενίσχυσης",
            year_summary["needs_support_count"]
        )

    with kpi7:
        st.metric(
            "Σχολές",
            year_summary["total_schools"]
        )

    with kpi8:
        st.metric(
            "Πόλεις",
            year_summary["total_cities"]
        )

    st.caption(
        "Η ένδειξη «100% στη ΓΕΛ Ημερήσια» αφορά την πλήρη κάλυψη της βασικής "
        "κατηγορίας εισαγωγής. Η ένδειξη «Με περιθώριο ενίσχυσης» αφορά "
        "προγράμματα που δεν κατέγραψαν πλήρη κάλυψη στη ΓΕΛ Ημερήσια."
    )

    st.divider()

    tab_text, tab_tables = st.tabs(
        [
            "Προεπισκόπηση δελτίου τύπου",
            "Υποστηρικτικοί πίνακες",
        ]
    )

    with tab_text:
        st.text_area(
            "Κείμενο δελτίου τύπου",
            value=press_release_text,
            height=700
        )

        word_file = create_word_press_release(
            text=press_release_text,
            selected_year=selected_year,
            top_base_display=top_base_display,
            top_admitted_display=top_admitted_display,
            top_base_change_display=top_base_change_display,
            gel_full_display=gel_full_display,
            support_display=support_display,
            full_display=full_display
        )

        st.download_button(
            label="⬇️ Λήψη δελτίου τύπου σε Word",
            data=word_file,
            file_name=f"dipae_press_release_{selected_year}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    with tab_tables:
        st.subheader("Top-5 υψηλότερες βάσεις ΓΕΛ Ημερήσια")

        st.dataframe(
            top_base_display,
            use_container_width=True,
            hide_index=True
        )

        st.subheader("Top-5 περισσότερων επιτυχόντων")

        st.dataframe(
            top_admitted_display,
            use_container_width=True,
            hide_index=True
        )

        if top_base_change_display is not None and not top_base_change_display.empty:
            st.subheader("Top-5 μεγαλύτερης αύξησης βάσης ΓΕΛ Ημερήσια")

            st.dataframe(
                top_base_change_display,
                use_container_width=True,
                hide_index=True
            )

        st.subheader("Προγράμματα με 100% κάλυψη στη ΓΕΛ Ημερήσια")

        st.dataframe(
            style_percent_table(gel_full_display),
            use_container_width=True,
            hide_index=True
        )

        st.subheader("Προγράμματα με περιθώριο ενίσχυσης")

        st.dataframe(
            support_display,
            use_container_width=True,
            hide_index=True
        )

        st.subheader("Αναλυτικός πίνακας προγραμμάτων")

        st.dataframe(
            style_percent_table(full_display),
            use_container_width=True,
            hide_index=True
        )

except Exception as e:
    st.error("Υπήρξε πρόβλημα κατά τη δημιουργία του δελτίου τύπου.")
    st.exception(e)