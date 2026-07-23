import io
from datetime import date

import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

from modules.database_manager import load_admissions_with_departments
from components.sidebar_branding import show_sidebar_branding


# ---------------------------------------------------------
# Ρυθμίσεις σελίδας
# ---------------------------------------------------------

st.set_page_config(
    page_title="Management Report | ΔΙΠΑΕ",
    page_icon="📄",
    layout="wide"
)

show_sidebar_branding()


st.title("📄 Διοικητική Αναφορά Εισακτέων ΔΙ.ΠΑ.Ε.")

st.markdown("""
Η σελίδα δημιουργεί αναφορά διοικητικής ενημέρωσης για τα στοιχεία εισακτέων
του ΔΙ.ΠΑ.Ε.

**Μεθοδολογικοί κανόνες:**

- Η συνολική εικόνα υπολογίζεται από **όλες τις κατηγορίες εισαγωγής**.
- Οι **Συνολικές Θέσεις** υπολογίζονται από τις **Αρχικές Θέσεις**.
- Η **Συνολική Κάλυψη** υπολογίζεται ως: Επιτυχόντες όλων των κατηγοριών / Συνολικές Θέσεις.
- Η **Κάλυψη ΓΕΛ Ημερήσια** υπολογίζεται μόνο για τη βασική κατηγορία ΓΕΛ Ημερήσια.
- Για τη ΓΕΛ Ημερήσια χρησιμοποιούνται οι θέσεις της συγκεκριμένης κατηγορίας.
- Η **Βάση Προγράμματος** είναι η **Βάση ΓΕΛ Ημερήσια**.
- Ο **Πρώτος Προγράμματος** είναι ο **Πρώτος ΓΕΛ Ημερήσια**.
- Δεν εμφανίζονται τελικές θέσεις ή φαινόμενη μεταβολή θέσεων.
- Δεν υπολογίζονται μέσοι όροι βάσεων.
""")

st.divider()


# ---------------------------------------------------------
# Βοηθητικές συναρτήσεις
# ---------------------------------------------------------

def safe_int(value):
    try:
        if pd.isna(value):
            return 0
        return int(round(float(value)))
    except Exception:
        return 0


def safe_float(value):
    try:
        if pd.isna(value):
            return 0.0
        return float(value)
    except Exception:
        return 0.0


def format_percent(value):
    try:
        return f"{float(value):.2f}%"
    except Exception:
        return "0.00%"


def format_score(value):
    try:
        if pd.isna(value):
            return ""
        return f"{float(value):.0f}"
    except Exception:
        return ""


def get_gel_day_data(df_year):
    """
    Επιστρέφει στοιχεία ΓΕΛ Ημερήσια ανά πρόγραμμα.

    Για τη ΓΕΛ Ημερήσια χρησιμοποιούνται οι final_positions,
    επειδή σε αυτή την κατηγορία αποτυπώνονται οι θέσεις κατόπιν μεταφοράς.
    """

    df_gel = df_year[
        df_year["exam_category"] == "ΓΕΛ Ημερήσια"
    ].copy()

    if df_gel.empty:
        return pd.DataFrame(
            columns=[
                "department_code",
                "gel_day_positions",
                "gel_day_admitted",
                "gel_day_empty",
                "gel_day_coverage",
                "gel_day_first_score",
                "gel_day_base_score",
            ]
        )

    df_gel["gel_day_positions"] = (
        df_gel["final_positions"]
        .fillna(df_gel["initial_positions"])
    )

    df_gel["gel_day_admitted"] = (
        df_gel["admitted"]
        .fillna(0)
    )

    df_gel["gel_day_empty"] = (
        df_gel["gel_day_positions"]
        - df_gel["gel_day_admitted"]
    )

    df_gel["gel_day_coverage"] = (
        df_gel["gel_day_admitted"]
        / df_gel["gel_day_positions"]
        * 100
    )

    gel_data = (
        df_gel[
            [
                "department_code",
                "gel_day_positions",
                "gel_day_admitted",
                "gel_day_empty",
                "gel_day_coverage",
                "first_score",
                "base_score",
            ]
        ]
        .drop_duplicates(subset=["department_code"])
        .rename(
            columns={
                "first_score": "gel_day_first_score",
                "base_score": "gel_day_base_score",
            }
        )
    )

    return gel_data


def build_department_summary(df_year):
    """
    Δημιουργεί σύνοψη ανά πρόγραμμα.
    """

    summary = (
        df_year
        .groupby(
            [
                "department_code",
                "department_name_clean",
                "school",
                "city",
            ],
            as_index=False
        )
        .agg(
            total_positions=("initial_positions", "sum"),
            total_admitted=("admitted", "sum"),
        )
    )

    summary["empty_positions"] = (
        summary["total_positions"]
        - summary["total_admitted"]
    )

    summary["total_coverage"] = (
        summary["total_admitted"]
        / summary["total_positions"]
        * 100
    )

    summary["total_coverage"] = (
        summary["total_coverage"]
        .fillna(0)
    )

    gel_data = get_gel_day_data(df_year)

    summary = summary.merge(
        gel_data,
        on="department_code",
        how="left"
    )

    for col in [
        "gel_day_positions",
        "gel_day_admitted",
        "gel_day_empty",
        "gel_day_coverage",
        "gel_day_first_score",
        "gel_day_base_score",
    ]:
        if col not in summary.columns:
            summary[col] = 0

        summary[col] = (
            summary[col]
            .fillna(0)
        )

    return summary


def classify_coverage(row):
    """
    Κατηγοριοποίηση κάλυψης προγράμματος.

    Η κατηγοριοποίηση βασίζεται στη συνολική κάλυψη και στη ΓΕΛ Ημερήσια.
    """

    total_coverage = safe_float(row.get("total_coverage", 0))
    gel_day_coverage = safe_float(row.get("gel_day_coverage", 0))
    gel_day_positions = safe_float(row.get("gel_day_positions", 0))

    if gel_day_positions <= 0:
        return "Χωρίς στοιχεία ΓΕΛ Ημερήσια"

    if total_coverage >= 99.999:
        return "Πλήρης συνολική κάλυψη"

    if gel_day_coverage >= 99.999 and total_coverage >= 95:
        return "Πλήρης κάλυψη ΓΕΛ Ημερήσια με μικρά κενά*"

    if gel_day_coverage >= 99.999:
        return "Πλήρης κάλυψη ΓΕΛ Ημερήσια με σημαντικά κενά*"

    return "Μη πλήρης κάλυψη ΓΕΛ Ημερήσια"


def build_year_summary(department_summary):
    """
    Δημιουργεί συνοπτικά στοιχεία έτους.
    """

    total_programs = int(department_summary["department_code"].nunique())
    total_schools = int(department_summary["school"].nunique())
    total_cities = int(department_summary["city"].nunique())

    total_positions = safe_int(
        department_summary["total_positions"].sum()
    )

    total_admitted = safe_int(
        department_summary["total_admitted"].sum()
    )

    total_empty = safe_int(
        department_summary["empty_positions"].sum()
    )

    total_coverage = (
        total_admitted / total_positions * 100
        if total_positions > 0
        else 0
    )

    gel_day_positions = safe_int(
        department_summary["gel_day_positions"]
        .fillna(0)
        .sum()
    )

    gel_day_admitted = safe_int(
        department_summary["gel_day_admitted"]
        .fillna(0)
        .sum()
    )

    gel_day_empty = safe_int(
        department_summary["gel_day_empty"]
        .fillna(0)
        .sum()
    )

    gel_day_coverage = (
        gel_day_admitted / gel_day_positions * 100
        if gel_day_positions > 0
        else 0
    )

    coverage_counts = (
        department_summary["coverage_category"]
        .value_counts()
        .to_dict()
    )

    return {
        "total_programs": total_programs,
        "total_schools": total_schools,
        "total_cities": total_cities,
        "total_positions": total_positions,
        "total_admitted": total_admitted,
        "total_empty": total_empty,
        "total_coverage": total_coverage,
        "gel_day_positions": gel_day_positions,
        "gel_day_admitted": gel_day_admitted,
        "gel_day_empty": gel_day_empty,
        "gel_day_coverage": gel_day_coverage,
        "full_total_count": coverage_counts.get("Πλήρης συνολική κάλυψη", 0),
        "gel_full_small_gaps_count": coverage_counts.get(
            "Πλήρης κάλυψη ΓΕΛ Ημερήσια με μικρά κενά*",
            0
        ),
        "gel_full_large_gaps_count": coverage_counts.get(
            "Πλήρης κάλυψη ΓΕΛ Ημερήσια με σημαντικά κενά*",
            0
        ),
        "gel_not_full_count": coverage_counts.get(
            "Μη πλήρης κάλυψη ΓΕΛ Ημερήσια",
            0
        ),
        "no_gel_data_count": coverage_counts.get(
            "Χωρίς στοιχεία ΓΕΛ Ημερήσια",
            0
        ),
    }


def build_school_summary(department_summary):
    """
    Σύνοψη ανά Σχολή.
    """

    school_summary = (
        department_summary
        .groupby("school", as_index=False)
        .agg(
            programs=("department_code", "nunique"),
            total_positions=("total_positions", "sum"),
            total_admitted=("total_admitted", "sum"),
            empty_positions=("empty_positions", "sum"),
            gel_day_positions=("gel_day_positions", "sum"),
            gel_day_admitted=("gel_day_admitted", "sum"),
            gel_day_empty=("gel_day_empty", "sum"),
        )
    )

    school_summary["total_coverage"] = (
        school_summary["total_admitted"]
        / school_summary["total_positions"]
        * 100
    )

    school_summary["gel_day_coverage"] = (
        school_summary["gel_day_admitted"]
        / school_summary["gel_day_positions"]
        * 100
    )

    school_summary["total_coverage"] = (
        school_summary["total_coverage"]
        .fillna(0)
    )

    school_summary["gel_day_coverage"] = (
        school_summary["gel_day_coverage"]
        .fillna(0)
    )

    return school_summary


def build_city_summary(department_summary):
    """
    Σύνοψη ανά Πόλη.
    """

    city_summary = (
        department_summary
        .groupby("city", as_index=False)
        .agg(
            programs=("department_code", "nunique"),
            total_positions=("total_positions", "sum"),
            total_admitted=("total_admitted", "sum"),
            empty_positions=("empty_positions", "sum"),
            gel_day_positions=("gel_day_positions", "sum"),
            gel_day_admitted=("gel_day_admitted", "sum"),
            gel_day_empty=("gel_day_empty", "sum"),
        )
    )

    city_summary["total_coverage"] = (
        city_summary["total_admitted"]
        / city_summary["total_positions"]
        * 100
    )

    city_summary["gel_day_coverage"] = (
        city_summary["gel_day_admitted"]
        / city_summary["gel_day_positions"]
        * 100
    )

    city_summary["total_coverage"] = (
        city_summary["total_coverage"]
        .fillna(0)
    )

    city_summary["gel_day_coverage"] = (
        city_summary["gel_day_coverage"]
        .fillna(0)
    )

    return city_summary


def format_program_table(df):
    """
    Αναλυτικός πίνακας προγραμμάτων για Word και προεπισκόπηση.
    """

    display = df[
        [
            "department_name_clean",
            "school",
            "city",
            "total_positions",
            "total_admitted",
            "empty_positions",
            "total_coverage",
            "gel_day_positions",
            "gel_day_admitted",
            "gel_day_empty",
            "gel_day_coverage",
            "gel_day_base_score",
            "gel_day_first_score",
            "coverage_category",
        ]
    ].copy()

    display = display.rename(
        columns={
            "department_name_clean": "Προπτυχιακό Πρόγραμμα",
            "school": "Σχολή",
            "city": "Πόλη",
            "total_positions": "Συνολικές Θέσεις",
            "total_admitted": "Επιτυχόντες",
            "empty_positions": "Κενές Θέσεις",
            "total_coverage": "Συνολική Κάλυψη %",
            "gel_day_positions": "ΓΕΛ Ημερ. Θέσεις",
            "gel_day_admitted": "ΓΕΛ Ημερ. Επιτυχόντες",
            "gel_day_empty": "Κενές ΓΕΛ Ημερ.",
            "gel_day_coverage": "Κάλυψη ΓΕΛ Ημ. %",
            "gel_day_base_score": "Βάση ΓΕΛ Ημ.",
            "gel_day_first_score": "Πρώτος ΓΕΛ Ημ.",
            "coverage_category": "Κατηγορία Κάλυψης",
        }
    )

    return clean_display_table(display)


def format_group_table(df, group_column, group_label):
    """
    Μορφοποίηση πίνακα ανά Σχολή ή Πόλη.
    """

    display = df[
        [
            group_column,
            "programs",
            "total_positions",
            "total_admitted",
            "empty_positions",
            "total_coverage",
            "gel_day_positions",
            "gel_day_admitted",
            "gel_day_empty",
            "gel_day_coverage",
        ]
    ].copy()

    display = display.rename(
        columns={
            group_column: group_label,
            "programs": "Προγράμματα",
            "total_positions": "Συνολικές Θέσεις",
            "total_admitted": "Επιτυχόντες",
            "empty_positions": "Κενές Θέσεις",
            "total_coverage": "Συνολική Κάλυψη %",
            "gel_day_positions": "ΓΕΛ Ημερ. Θέσεις",
            "gel_day_admitted": "ΓΕΛ Ημερ. Επιτυχόντες",
            "gel_day_empty": "Κενές ΓΕΛ Ημερ.",
            "gel_day_coverage": "Κάλυψη ΓΕΛ Ημ. %",
        }
    )

    return clean_display_table(display)


def clean_display_table(display):
    """
    Καθαρίζει τύπους και μορφοποιεί αριθμούς.
    """

    integer_columns = [
        "Προγράμματα",
        "Συνολικές Θέσεις",
        "Επιτυχόντες",
        "Κενές Θέσεις",
        "ΓΕΛ Ημερ. Θέσεις",
        "ΓΕΛ Ημερ. Επιτυχόντες",
        "Κενές ΓΕΛ Ημερ.",
        "Βάση ΓΕΛ Ημ.",
        "Πρώτος ΓΕΛ Ημ.",
    ]

    for col in integer_columns:
        if col in display.columns:
            display[col] = (
                display[col]
                .fillna(0)
                .round(0)
                .astype(int)
            )

    percent_columns = [
        "Συνολική Κάλυψη %",
        "Κάλυψη ΓΕΛ Ημ. %",
    ]

    for col in percent_columns:
        if col in display.columns:
            display[col] = (
                display[col]
                .fillna(0)
                .round(2)
            )

    return display


def style_table(df):
    """
    Styling προεπισκόπησης.
    """

    style_obj = df.style

    format_dict = {}

    for col in df.columns:
        if "Κάλυψη" in col and "%" in col:
            format_dict[col] = "{:.2f}%"

    if format_dict:
        style_obj = style_obj.format(format_dict)

    return style_obj


# ---------------------------------------------------------
# Δημιουργία Word
# ---------------------------------------------------------

def add_heading(document, text, level=1):
    paragraph = document.add_heading(text, level=level)
    return paragraph


def add_paragraph(document, text, bold=False):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return paragraph


def add_metric_table(document, year_summary):
    """
    Προσθέτει συνοπτικό πίνακα δεικτών στο Word.
    """

    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = [
        "Συνολικές Θέσεις",
        "Επιτυχόντες",
        "Κενές Θέσεις",
        "Συνολική Κάλυψη",
    ]

    values = [
        str(year_summary["total_positions"]),
        str(year_summary["total_admitted"]),
        str(year_summary["total_empty"]),
        format_percent(year_summary["total_coverage"]),
    ]

    for i, cell in enumerate(table.rows[0].cells):
        cell.text = f"{headers[i]}\n{values[i]}"
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    document.add_paragraph()

    table_gel = document.add_table(rows=1, cols=4)
    table_gel.style = "Table Grid"
    table_gel.alignment = WD_TABLE_ALIGNMENT.CENTER

    gel_headers = [
        "ΓΕΛ Ημερ. Θέσεις",
        "ΓΕΛ Ημερ. Επιτυχόντες",
        "Κενές ΓΕΛ Ημερ.",
        "Κάλυψη ΓΕΛ Ημερ.",
    ]

    gel_values = [
        str(year_summary["gel_day_positions"]),
        str(year_summary["gel_day_admitted"]),
        str(year_summary["gel_day_empty"]),
        format_percent(year_summary["gel_day_coverage"]),
    ]

    for i, cell in enumerate(table_gel.rows[0].cells):
        cell.text = f"{gel_headers[i]}\n{gel_values[i]}"
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    document.add_paragraph()


def add_dataframe_table(document, df, title=None):
    """
    Προσθέτει DataFrame ως πίνακα στο Word.
    """

    if title:
        add_heading(document, title, level=2)

    if df.empty:
        add_paragraph(document, "Δεν υπάρχουν διαθέσιμα δεδομένα.")
        return

    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells

    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)

    for _, row in df.iterrows():
        row_cells = table.add_row().cells

        for i, col in enumerate(df.columns):
            value = row[col]

            if pd.isna(value):
                text = ""
            elif isinstance(value, float):
                if "Κάλυψη" in str(col) and "%" in str(col):
                    text = f"{value:.2f}%"
                else:
                    text = f"{value:.2f}"
            else:
                text = str(value)

            row_cells[i].text = text

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    document.add_paragraph()


def create_word_report(
    selected_year,
    year_summary,
    department_summary,
    top_base,
    top_admitted,
    top_empty,
    low_coverage,
    school_display,
    city_display,
    full_display,
):
    """
    Δημιουργεί Word διοικητικής αναφοράς.
    """

    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)

    title = document.add_heading(
        f"Διοικητική Αναφορά Εισακτέων ΔΙ.ΠΑ.Ε. {selected_year}",
        level=0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph(
        f"Ημερομηνία δημιουργίας: {date.today().strftime('%d/%m/%Y')}"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph()

    add_heading(document, "1. Συνοπτική εικόνα Ιδρύματος", level=1)

    add_metric_table(document, year_summary)

    add_paragraph(
        document,
        (
            f"Το ΔΙ.ΠΑ.Ε. διαθέτει {year_summary['total_programs']} ενεργά προπτυχιακά "
            f"προγράμματα σπουδών, κατανεμημένα σε {year_summary['total_schools']} σχολές "
            f"και {year_summary['total_cities']} πόλεις."
        )
    )

    add_paragraph(
        document,
        (
            f"Οι συνολικές θέσεις όλων των κατηγοριών ανέρχονται σε "
            f"{year_summary['total_positions']}, ενώ οι επιτυχόντες σε "
            f"{year_summary['total_admitted']}. Η συνολική κάλυψη διαμορφώνεται σε "
            f"{format_percent(year_summary['total_coverage'])}."
        )
    )

    add_paragraph(
        document,
        (
            f"Στη ΓΕΛ Ημερήσια οι θέσεις ανέρχονται σε "
            f"{year_summary['gel_day_positions']}, οι επιτυχόντες σε "
            f"{year_summary['gel_day_admitted']} και η κάλυψη ΓΕΛ Ημερήσια σε "
            f"{format_percent(year_summary['gel_day_coverage'])}."
        )
    )

    add_heading(document, "2. Κατηγορίες κάλυψης προγραμμάτων", level=1)

    add_paragraph(
        document,
        (
            f"Πλήρη συνολική κάλυψη εμφανίζουν {year_summary['full_total_count']} προγράμματα. "
            f"Πλήρη κάλυψη στη ΓΕΛ Ημερήσια με μικρά κενά εμφανίζουν "
            f"{year_summary['gel_full_small_gaps_count']} προγράμματα, ενώ πλήρη κάλυψη στη ΓΕΛ "
            f"Ημερήσια με σημαντικά κενά εμφανίζουν "
            f"{year_summary['gel_full_large_gaps_count']} προγράμματα. "
            f"Μη πλήρη κάλυψη στη ΓΕΛ Ημερήσια εμφανίζουν "
            f"{year_summary['gel_not_full_count']} προγράμματα."
        )
    )

    add_paragraph(
        document,
        (
            "* Ως μικρά κενά θεωρείται συνολική κάλυψη τουλάχιστον 95% "
            "αλλά μικρότερη από 100%, με πλήρη κάλυψη στη ΓΕΛ Ημερήσια."
        )
    )

    add_dataframe_table(document, top_base, "3. Top-5 υψηλότερων βάσεων ΓΕΛ Ημερήσια")
    add_dataframe_table(document, top_admitted, "4. Top-5 περισσότερων επιτυχόντων")
    add_dataframe_table(document, top_empty, "5. Top-5 περισσότερων κενών θέσεων")
    add_dataframe_table(document, low_coverage, "6. Top-5 χαμηλότερης συνολικής κάλυψης")
    add_dataframe_table(document, school_display, "7. Ανάλυση ανά Σχολή")
    add_dataframe_table(document, city_display, "8. Ανάλυση ανά Πόλη")
    add_dataframe_table(document, full_display, "9. Αναλυτικός πίνακας προγραμμάτων")

    add_heading(document, "10. Μεθοδολογικές σημειώσεις", level=1)

    add_paragraph(
        document,
        (
            "Η Συνολική Κάλυψη υπολογίζεται επί των συνολικών θέσεων όλων των κατηγοριών "
            "εισαγωγής. Οι Συνολικές Θέσεις προκύπτουν από τις αρχικές θέσεις."
        )
    )

    add_paragraph(
        document,
        (
            "Η Κάλυψη ΓΕΛ Ημερήσια υπολογίζεται αποκλειστικά επί των θέσεων της ΓΕΛ Ημερήσιας. "
            "Για τη συγκεκριμένη κατηγορία χρησιμοποιούνται οι θέσεις της κατηγορίας, όπως "
            "αποτυπώνονται στα επίσημα δεδομένα."
        )
    )

    add_paragraph(
        document,
        (
            "Η Βάση Προγράμματος και ο Πρώτος Προγράμματος αντιστοιχούν στη Βάση και στον "
            "Πρώτο της ΓΕΛ Ημερήσιας."
        )
    )

    output = io.BytesIO()
    document.save(output)
    output.seek(0)

    return output


# ---------------------------------------------------------
# Κύρια ροή
# ---------------------------------------------------------

try:
    df = load_admissions_with_departments()

    if df.empty:
        st.warning("Δεν υπάρχουν ακόμη δεδομένα εισακτέων στη βάση.")
        st.stop()

    years = sorted(df["year"].dropna().unique().tolist())

    selected_year = st.selectbox(
        "Έτος",
        years,
        index=len(years) - 1
    )

    df_year = df[df["year"] == selected_year].copy()

    if df_year.empty:
        st.warning("Δεν υπάρχουν δεδομένα για το επιλεγμένο έτος.")
        st.stop()

    department_summary = build_department_summary(df_year)

    department_summary["coverage_category"] = department_summary.apply(
        classify_coverage,
        axis=1
    )

    year_summary = build_year_summary(department_summary)

    # ---------------------------------------------------------
    # Πίνακες αναφοράς
    # ---------------------------------------------------------

    top_base = format_program_table(
        department_summary[
            department_summary["gel_day_base_score"] > 0
        ]
        .sort_values(
            "gel_day_base_score",
            ascending=False
        )
        .head(5)
    )

    top_admitted = format_program_table(
        department_summary
        .sort_values(
            "total_admitted",
            ascending=False
        )
        .head(5)
    )

    top_empty = format_program_table(
        department_summary
        .sort_values(
            "empty_positions",
            ascending=False
        )
        .head(5)
    )

    low_coverage = format_program_table(
        department_summary
        .sort_values(
            "total_coverage",
            ascending=True
        )
        .head(5)
    )

    school_summary = build_school_summary(department_summary)

    school_display = format_group_table(
        school_summary.sort_values(
            "total_coverage",
            ascending=True
        ),
        group_column="school",
        group_label="Σχολή"
    )

    city_summary = build_city_summary(department_summary)

    city_display = format_group_table(
        city_summary.sort_values(
            "total_coverage",
            ascending=True
        ),
        group_column="city",
        group_label="Πόλη"
    )

    full_display = format_program_table(
        department_summary.sort_values(
            "gel_day_base_score",
            ascending=False
        )
    )

    # ---------------------------------------------------------
    # Προεπισκόπηση
    # ---------------------------------------------------------

    st.subheader("Προεπισκόπηση αναφοράς")

    # 1η σειρά: Γενική εικόνα όλων των κατηγοριών
    preview_kpi1, preview_kpi2, preview_kpi3, preview_kpi4 = st.columns(4)

    with preview_kpi1:
        st.metric(
            "Συνολικές Θέσεις",
            year_summary["total_positions"]
        )

    with preview_kpi2:
        st.metric(
            "Επιτυχόντες",
            year_summary["total_admitted"]
        )

    with preview_kpi3:
        st.metric(
            "Κενές Θέσεις",
            year_summary["total_empty"]
        )

    with preview_kpi4:
        st.metric(
            "Συνολική Κάλυψη",
            f'{year_summary["total_coverage"]:.2f}%'
        )

    # 2η σειρά: ΓΕΛ Ημερήσια
    gel_kpi1, gel_kpi2, gel_kpi3, gel_kpi4 = st.columns(4)

    with gel_kpi1:
        st.metric(
            "ΓΕΛ Ημερ. Θέσεις",
            year_summary["gel_day_positions"]
        )

    with gel_kpi2:
        st.metric(
            "ΓΕΛ Ημερ. Επιτυχόντες",
            year_summary["gel_day_admitted"]
        )

    with gel_kpi3:
        st.metric(
            "Κενές ΓΕΛ Ημερ.",
            year_summary["gel_day_empty"]
        )

    with gel_kpi4:
        st.metric(
            "Κάλυψη ΓΕΛ Ημερ.",
            f'{year_summary["gel_day_coverage"]:.2f}%'
        )

    st.caption(
        "Η πρώτη σειρά αφορά όλες τις κατηγορίες εισαγωγής. "
        "Η δεύτερη σειρά αφορά αποκλειστικά τη ΓΕΛ Ημερήσια."
    )

    st.divider()

    st.subheader("Πίνακες που θα συμπεριληφθούν στην αναφορά")

    tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs(
        [
            "Top Βάσεις",
            "Top Επιτυχόντες",
            "Top Κενά",
            "Χαμηλή Κάλυψη",
            "Σχολές / Πόλεις",
            "Πλήρης Πίνακας",
        ]
    )

    with tab1:
        st.dataframe(
            style_table(top_base),
            use_container_width=True,
            hide_index=True
        )

    with tab2:
        st.dataframe(
            style_table(top_admitted),
            use_container_width=True,
            hide_index=True
        )

    with tab3:
        st.dataframe(
            style_table(top_empty),
            use_container_width=True,
            hide_index=True
        )

    with tab4:
        st.dataframe(
            style_table(low_coverage),
            use_container_width=True,
            hide_index=True
        )

    with tab5:
        st.markdown("#### Ανάλυση ανά Σχολή")
        st.dataframe(
            style_table(school_display),
            use_container_width=True,
            hide_index=True
        )

        st.markdown("#### Ανάλυση ανά Πόλη")
        st.dataframe(
            style_table(city_display),
            use_container_width=True,
            hide_index=True
        )

    with tab6:
        st.dataframe(
            style_table(full_display),
            use_container_width=True,
            hide_index=True
        )

    st.divider()

    # ---------------------------------------------------------
    # Δημιουργία Word
    # ---------------------------------------------------------

    word_file = create_word_report(
        selected_year=selected_year,
        year_summary=year_summary,
        department_summary=department_summary,
        top_base=top_base,
        top_admitted=top_admitted,
        top_empty=top_empty,
        low_coverage=low_coverage,
        school_display=school_display,
        city_display=city_display,
        full_display=full_display,
    )

    st.download_button(
        label="⬇️ Κατέβασμα διοικητικής αναφοράς Word",
        data=word_file,
        file_name=f"management_report_dipae_{selected_year}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    st.info(
        "Το Word περιλαμβάνει τη συνοπτική εικόνα Ιδρύματος, τους βασικούς πίνακες "
        "και τις μεθοδολογικές σημειώσεις."
    )

except Exception as e:
    st.error("Υπήρξε πρόβλημα κατά τη δημιουργία της διοικητικής αναφοράς.")
    st.exception(e)